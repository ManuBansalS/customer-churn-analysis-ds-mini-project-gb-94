"""
Convert PROJECT_DOCUMENTATION.md to Word document (.docx)
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set borders on table cells"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def markdown_to_docx(md_file, docx_file):
    """Convert markdown to Word document"""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Split by lines
    lines = content.split('\n')
    
    i = 0
    current_code_block = []
    in_code_block = False
    current_table = None
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            
        # Heading 2
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            i += 1
            
        # Heading 3
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            i += 1
            
        # Heading 4
        elif line.startswith('#### '):
            title = line[5:].strip()
            doc.add_heading(title, level=4)
            i += 1
            
        # Horizontal rule
        elif line.strip().startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            
        # Code block
        elif line.strip().startswith('```'):
            # Collect code block
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_block.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Add code to document
            if code_block:
                code_text = '\n'.join(code_block)
                code_para = doc.add_paragraph(code_text, style='No Spacing')
                code_para.style = 'List Bullet'
                
                # Format code
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E8E8E8')
                code_para._element.get_or_add_pPr().append(shading_elm)
                
        # Table
        elif line.strip().startswith('|'):
            # Collect table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [cell.strip() for cell in lines[i].split('|')][1:-1]
                table_rows.append(cells)
                i += 1
            
            if table_rows:
                # Determine number of columns
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        
                        # Header row formatting
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), '366092')
                            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            
        # Numbered list
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s(.*)', line.strip())
            if match:
                text = match.group(2)
                doc.add_paragraph(text, style='List Number')
            i += 1
            
        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)        # Code
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
            
            if text:
                doc.add_paragraph(text)
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Document created: {docx_file}")

if __name__ == '__main__':
    md_file = 'PROJECT_DOCUMENTATION.md'
    docx_file = 'PROJECT_DOCUMENTATION.docx'
    
    markdown_to_docx(md_file, docx_file)
